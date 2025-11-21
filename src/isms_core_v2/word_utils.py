"""
word_utils.py

Helpers for adding content under existing headings in a Word template using
python-docx, now with support for:

- preserving inline bold/italic/underline
- rendering hyperlinks (blue + underlined) from run metadata

Note: Custom Properties inside DOCX are not edited here (requires additional libs).
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional, Sequence

from docx.document import Document as _Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _find_heading_paragraph(doc: _Document, heading_text: str) -> Optional[Paragraph]:
    """Return the first paragraph whose text matches heading_text (case-insensitive, trimmed)."""
    target = (heading_text or "").strip().lower()
    if not target:
        return None

    for p in doc.paragraphs:
        if p.text.strip().lower() == target:
            return p
    return None


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """
    Insert and return a new paragraph immediately *after* the given one.

    python-docx doesn't expose `insert_after`, so we do a small XML shuffle.
    """
    new_p = paragraph.insert_paragraph_before("")
    # swap text: insert_paragraph_before puts it above; we want effectively "after"
    new_p.text, paragraph.text = paragraph.text, new_p.text
    return new_p


# ---------------------------------------------------------------------------
# Hyperlink + rich run rendering
# ---------------------------------------------------------------------------

def add_hyperlink_run(
    paragraph: Paragraph,
    url: str,
    text: str,
    bold: bool = False,
    italic: bool = False,
    underline: bool = True,
) -> None:
    """
    Add a single hyperlink run to `paragraph` with Word-style formatting:

    - blue font colour
    - underline (by default)
    - optional bold/italic flags

    This uses low-level XML because python-docx doesn't expose hyperlink creation
    directly for runs.
    """
    part = paragraph.part
    # Create relationship ID for hyperlink
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create the run
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # Bold
    if bold:
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "true")
        r_pr.append(b)

    # Italic
    if italic:
        i = OxmlElement("w:i")
        i.set(qn("w:val"), "true")
        r_pr.append(i)

    # Underline
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    new_run.append(r_pr)

    # Text node
    t = OxmlElement("w:t")
    t.text = text or ""
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_rich_paragraph(
    doc: _Document,
    runs: Sequence[Dict[str, Any]],
    style_name: str = "ISMS Body",
) -> Paragraph:
    """
    Render a paragraph composed of multiple run fragments, each containing:

        {
          "text": str,
          "bold": bool,
          "italic": bool,
          "underline": bool,
          "hyperlink": str | None
        }

    If `hyperlink` is not None, a hyperlink run is created; otherwise a normal run.
    """
    p = doc.add_paragraph()
    # Apply style (fallback to Normal)
    for s in (style_name, "Normal"):
        try:
            p.style = s
            break
        except (KeyError, ValueError):
            continue

    for frag in runs:
        text = frag.get("text") or ""
        if not text:
            continue

        bold = bool(frag.get("bold"))
        italic = bool(frag.get("italic"))
        underline = bool(frag.get("underline"))
        href = frag.get("hyperlink")

        if href:
            # Hyperlink run
            add_hyperlink_run(
                paragraph=p,
                url=str(href),
                text=text,
                bold=bold,
                italic=italic,
                underline=underline or True,
            )
        else:
            # Normal run
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline

    return p


# ---------------------------------------------------------------------------
# Heading-aware helpers (backwards-compatible)
# ---------------------------------------------------------------------------

def add_body_under_heading(
    doc: _Document,
    heading_text: str,
    body: Any,
) -> None:
    """
    Append body content under a heading, creating the heading if needed with ISMS styles.

    `body` can be:
      - str: treated as plain text and split into paragraphs on line breaks (old behaviour)
      - list[dict]: treated as a list of run fragments for a *single* rich paragraph
                    (e.g. the `runs` array from JSON)
    """
    p = _find_heading_paragraph(doc, heading_text)
    if p is None:
        # Create a new heading with ISMS Heading 1 if available
        p = doc.add_paragraph(heading_text or "")
        for style_name in ("ISMS Heading 1", "Heading 1"):
            try:
                p.style = style_name
                break
            except KeyError:
                continue

    # Insert content *after* the heading
    if isinstance(body, str) or body is None:
        text = str(body or "")
        for line in text.splitlines():
            if not line.strip():
                continue
            new_p = _insert_paragraph_after(p)
            new_p.text = line.strip()
            for style_name in ("ISMS Body", "Normal"):
                try:
                    new_p.style = style_name
                    break
                except KeyError:
                    continue
            p = new_p  # so subsequent lines are under the last inserted paragraph

    elif isinstance(body, list) and body and isinstance(body[0], dict) and "text" in body[0]:
        # Assume `body` is a list of run fragments (one paragraph)
        # We'll create that paragraph under the heading.
        new_p = _insert_paragraph_after(p)
        # Temporarily create another paragraph, render runs there, then move XML
        tmp_doc = doc  # same doc
        rich_p = render_rich_paragraph(tmp_doc, body)
        # Move runs from rich_p to new_p
        for r in list(rich_p._p):
            new_p._p.append(r)
        # Remove the temporary empty paragraph
        tmp_doc._body._body.remove(rich_p._p)
    else:
        # Fallback: stringified
        text = str(body)
        new_p = _insert_paragraph_after(p)
        new_p.text = text
        for style_name in ("ISMS Body", "Normal"):
            try:
                new_p.style = style_name
                break
            except KeyError:
                continue


def add_rich_blocks_under_heading(
    doc: _Document,
    heading_text: str,
    blocks: Sequence[Dict[str, Any]],
) -> None:
    """
    Convenience helper: given a list of content blocks (e.g. from JSON section.content),
    render them under a heading.

    - If block["type"] == "paragraph" and "runs" present -> use rich rendering.
    - Else if block["type"] == "paragraph" and only "text" present -> plain text.
    - Other block types can be added as needed (tables, lists, etc.).
    """
    p = _find_heading_paragraph(doc, heading_text)
    if p is None:
        # Create the heading if needed
        p = doc.add_paragraph(heading_text or "")
        for style_name in ("ISMS Heading 1", "Heading 1"):
            try:
                p.style = style_name
                break
            except KeyError:
                continue

    for block in blocks:
        btype = block.get("type")
        if btype == "paragraph":
            runs = block.get("runs")
            text = block.get("text", "")

            if runs:
                rich_p = render_rich_paragraph(doc, runs)
                # Ensure it's positioned after the heading / previous content
                p = rich_p
            else:
                # Plain paragraph
                new_p = _insert_paragraph_after(p)
                new_p.text = text
                for style_name in ("ISMS Body", "Normal"):
                    try:
                        new_p.style = style_name
                        break
                    except KeyError:
                        continue
                p = new_p

        # TODO: handle tables, lists, etc. as needed


__all__ = [
    "_find_heading_paragraph",
    "add_body_under_heading",
    "add_hyperlink_run",
    "render_rich_paragraph",
    "add_rich_blocks_under_heading",
]
