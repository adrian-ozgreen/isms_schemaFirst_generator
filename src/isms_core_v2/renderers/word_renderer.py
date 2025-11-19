"""
word_renderer.py

Schema-driven renderer:
- Loads a base ISMS Word template
- Updates core properties and Document Control table from DocMetadata
- Appends sections & subsections with appropriate ISMS Heading styles
- Renders simple content blocks (paragraph, bullet_list, numbered_list)
- Preserves existing Title Page and TOC layout
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Dict, Callable, Set

from docx import Document
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document
from docx.table import _Cell, Table

from ..models import DocumentModel, Section, ContentBlock, DocMetadata  # adjust import if needed

# Placeholder → metadata resolver (all supported placeholders)
PLACEHOLDER_MAP: Dict[str, Callable[[DocMetadata], str]] = {
    "[[DOC_ID]]":               lambda m: m.doc_id,
    "[[DOC_TITLE]]":            lambda m: m.title,
    "[[DOC_VERSION]]":          lambda m: m.version,
    "[[DOC_OWNER]]":            lambda m: m.owner,
    "[[DOC_APPROVER]]":         lambda m: m.approver or "",
    "[[DOC_STATUS]]":           lambda m: m.status,
    "[[DOC_TYPE]]":             lambda m: m.doc_type,
    "[[DOC_CONFIDENTIALITY]]":  lambda m: m.confidentiality or "",
    "[[DOC_DATE_COMPLETED]]":   lambda m: m.date_completed or "",
    "[[DOC_NEXT_REVIEW_DATE]]": lambda m: m.next_review_date or "",
}

# We can treat *all* of them as required in the template if you want:
REQUIRED_PLACEHOLDERS: Set[str] = set(PLACEHOLDER_MAP.keys())


def _apply_metadata_placeholders(doc: _Document, metadata: DocMetadata) -> None:
    """
    Replace all supported [[PLACEHOLDER]] tokens in:
    - body paragraphs
    - tables
    - headers
    - footers
    with values from DocMetadata.
    """

    # Resolve all placeholder values once
    resolved_map = {
        placeholder: resolver(metadata)
        for placeholder, resolver in PLACEHOLDER_MAP.items()
    }

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            original = p.text or ""
            if not original:
                continue
            new_text = original
            for placeholder, value in resolved_map.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, value)
            if new_text != original:
                # This resets runs, which is fine for metadata-only paragraphs.
                p.text = new_text

    # Document body
    replace_in_paragraphs(doc.paragraphs)

    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Headers & footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        replace_in_paragraphs(header.paragraphs)
        replace_in_paragraphs(footer.paragraphs)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)





# -------------------------------------------------------------------
# Style maps and helpers
# -------------------------------------------------------------------

# Map section.level → candidate heading styles (first one that exists is used)
HEADING_STYLE_MAP: dict[int, tuple[str, ...]] = {
    1: ("ISMS Heading 1", "Heading 1"),
    2: ("ISMS Heading 2", "Heading 2"),
    3: ("ISMS Heading 3", "Heading 3"),
    4: ("ISMS Heading 4", "Heading 4"),
}

# Body and list style fallbacks
BODY_STYLE_CANDIDATES: tuple[str, ...] = ("ISMS Body", "Normal")
BULLET_STYLE_CANDIDATES: tuple[str, ...] = (
    "ISMS List Bullet",
    "List Bullet",
    "List Paragraph",
    "Normal",
)
NUMBERED_STYLE_CANDIDATES: tuple[str, ...] = (
    "ISMS List Numbered",
    "List Number",
    "List Paragraph",
    "Normal",
)

# Sections whose layout is assumed to be baked into the template
RESERVED_TEMPLATE_SECTIONS = {
    "title_page",
    "document_control",
    "table_of_contents",
}


def _apply_first_existing_style(paragraph: Paragraph, candidates: Iterable[str]) -> None:
    """
    Try each style name in order; apply the first one that exists in the document.
    """
    for style_name in candidates:
        try:
            paragraph.style = style_name
            return
        except KeyError:
            continue
    # If none found, leave default style


def _apply_first_existing_table_style(table: Table, candidates: Iterable[str]) -> None:
    """
    Try each table style name in order; apply the first one that exists.
    """
    for style_name in candidates:
        try:
            table.style = style_name
            return
        except KeyError:
            continue
    # If none found, leave Word's default table style


def _render_table_block(doc: Document, block: ContentBlock) -> None:
    """
    Render a ContentBlock of kind='table' as a Word table, using the
    'TracWater table' style if available.
    """

    header = block.header or []
    rows = block.rows or []

    # Determine table dimensions
    num_header_cols = len(header)
    max_row_cols = max((len(r) for r in rows), default=0)
    cols = max(num_header_cols, max_row_cols)

    if cols == 0:
        # Nothing meaningful to render
        return

    # Create table: +1 row if header present
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=cols)

    # Apply table style with fallbacks
    _apply_first_existing_table_style(
        table,
        ("TracWater table", "ISMS Table", "Table Grid")
    )

    current_row = 0

    # Fill header row if present
    if header:
        hdr_cells = table.rows[current_row].cells
        for idx, value in enumerate(header):
            if idx >= cols:
                break
            cell = hdr_cells[idx]
            cell.text = value
            # Apply paragraph style if defined
            for p in cell.paragraphs:
                _apply_first_existing_style(p, ("ISMS Body", "Normal"))
        current_row += 1

    # Fill body rows
    for r_values in rows:
        row_cells = table.rows[current_row].cells
        for idx, value in enumerate(r_values):
            if idx >= cols:
                break
            cell = row_cells[idx]
            cell.text = value
            for p in cell.paragraphs:
                _apply_first_existing_style(p, ("ISMS Body", "Normal"))
        current_row += 1

    # Optional caption (below table)
    if block.caption:
        caption_para = doc.add_paragraph(block.caption)
        _apply_first_existing_style(caption_para, ("ISMS Body", "Normal"))




# -------------------------------------------------------------------
# Metadata → document properties / control table
# -------------------------------------------------------------------

def _update_core_properties(doc: Document, metadata: DocMetadata) -> None:
    """
    Map DocMetadata into Word core properties.
    (These are visible in File → Info and some fields on title page / headers.)
    """
    core = doc.core_properties
    core.title = metadata.title
    core.subject = metadata.doc_type
    core.category = metadata.status
    core.creator = metadata.owner
    # Optional extras:
    core.keywords = f"{metadata.doc_id};{metadata.doc_type}"
    # core.comments = "Generated by ISMS Hybrid v2"


# Normalisation helper for table labels
def _normalise_label(text: str) -> str:
    return (
        text.strip()
        .lower()
        .replace(" ", "")
        .replace("\n", "")
        .replace(":", "")
    )


# Labels we recognise in the first cell of a row → which metadata field to write
DOC_CONTROL_LABEL_ALIASES: dict[str, set[str]] = {
    "doc_id": {"docid", "docid#", "documentid", "documentid#", "docidno", "documentidno"},
    "version": {"version", "rev", "revision"},
    "owner": {"owner", "documentowner", "docowner", "documentownername"},
    "status": {"status", "documentstatus", "docstatus", "approvalstatus"},
}


def _get_metadata_value(metadata: DocMetadata, key: str) -> str:
    if key == "doc_id":
        return metadata.doc_id
    if key == "version":
        return metadata.version
    if key == "owner":
        return metadata.owner
    if key == "status":
        return metadata.status
    return ""


def _update_document_control_table(doc: Document, metadata: DocMetadata) -> None:
    """
    Find the Document Control table and write DocID, Version, Owner, Status
    based on label matching in the first cell of each row.

    Assumptions:
    - There is a table on the Document Control page whose rows contain labels like:
        "Doc ID", "Document ID", "Version", "Owner", "Status"
    - The value is in the second cell of the row (or last cell if only 1).
    """
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue

            label_cell = row.cells[0]
            label_norm = _normalise_label(label_cell.text)

            for meta_key, aliases in DOC_CONTROL_LABEL_ALIASES.items():
                if label_norm in aliases:
                    # choose value cell: second cell if exists, else last
                    value_cell = row.cells[1] if len(row.cells) > 1 else row.cells[-1]
                    value_cell.text = _get_metadata_value(metadata, meta_key)
                    break
        # We don't break out of table loop because other tables may also exist;
        # this is safe and idempotent.


# -------------------------------------------------------------------
# Content rendering
# -------------------------------------------------------------------

def _add_paragraph_block(doc: Document, block: ContentBlock) -> None:
    p = doc.add_paragraph(block.text if isinstance(block.text, str) else "")
    _apply_first_existing_style(p, BODY_STYLE_CANDIDATES)


def _add_bullet_list_block(doc: Document, block: ContentBlock) -> None:
    items = block.text if isinstance(block.text, list) else [str(block.text)]
    for text in items:
        p = doc.add_paragraph(text or "")
        _apply_first_existing_style(p, BULLET_STYLE_CANDIDATES)


def _add_numbered_list_block(doc: Document, block: ContentBlock) -> None:
    items = block.text if isinstance(block.text, list) else [str(block.text)]
    for text in items:
        p = doc.add_paragraph(text or "")
        _apply_first_existing_style(p, NUMBERED_STYLE_CANDIDATES)


# def _render_content_block(doc: Document, block: ContentBlock) -> None:
#     if block.kind == "paragraph":
#         _add_paragraph_block(doc, block)
#     elif block.kind == "bullet_list":
#         _add_bullet_list_block(doc, block)
#     elif block.kind == "numbered_list":
#         _add_numbered_list_block(doc, block)
#     else:
#         # For now we ignore unknown kinds; could log a warning later
#         _add_paragraph_block(doc, block)


def _render_content_block(doc: Document, block: ContentBlock) -> None:
    if block.kind == "paragraph":
        for line in str(block.text or "").splitlines():
            if not line.strip():
                continue
            p = doc.add_paragraph(line.strip())
            _apply_first_existing_style(p, ("ISMS Body", "Normal"))

    elif block.kind == "bullet_list":
        for item in block.text or []:
            p = doc.add_paragraph(str(item or ""))
            _apply_first_existing_style(p, BULLET_STYLE_CANDIDATES)

    elif block.kind == "numbered_list":
        for item in block.text or []:
            p = doc.add_paragraph(str(item or ""))
            _apply_first_existing_style(p, NUMBERED_STYLE_CANDIDATES)

    elif block.kind == "table":
        _render_table_block(doc, block)

    else:
        # Future-proof: ignore unknown kinds gracefully
        return





# -------------------------------------------------------------------
# Section rendering
# -------------------------------------------------------------------

def _add_section_heading(doc: Document, section: Section) -> Paragraph:
    """
    Append a section heading at the end of the document with an appropriate ISMS Heading style.
    """
    heading_paragraph = doc.add_paragraph(section.title or "")
    style_candidates = HEADING_STYLE_MAP.get(section.level, ("Heading 1",))
    _apply_first_existing_style(heading_paragraph, style_candidates)
    return heading_paragraph


def _render_section_recursive(doc: Document, section: Section) -> None:
    """
    Render a section heading + content + all subsections, appended at the end of the document.
    """
    # Append heading
    _add_section_heading(doc, section)

    # Append content blocks
    for block in section.content:
        _render_content_block(doc, block)

    # Render subsections after content
    for sub in section.subsections:
        _render_section_recursive(doc, sub)


# -------------------------------------------------------------------
# Reserved sections: title page, document control, TOC
# -------------------------------------------------------------------

def _render_title_page(doc: Document, model: DocumentModel, section: Section) -> None:
    """
    Reserved hook for future mapping of metadata/title into the template's title page.

    For now, we only update core properties; the visual layout stays as-is.
    """
    _update_core_properties(doc, model.metadata)


def _render_document_control(doc: Document, model: DocumentModel, section: Section) -> None:
    """
    Map DocMetadata into the Document Control table (if found).
    """
    _update_document_control_table(doc, model.metadata)


def _render_table_of_contents(doc: Document, model: DocumentModel, section: Section) -> None:
    """
    Reserved for future TOC field updates.

    python-docx cannot update TOC fields directly; typically, Word users will update fields
    (F9) after generation. For now, this is a NO-OP.
    """
    return


def _dispatch_reserved_section(doc: Document, model: DocumentModel, section: Section) -> None:
    """
    Call the appropriate handler for sections whose layout is pre-built in the template.
    """
    if section.key == "title_page":
        _render_title_page(doc, model, section)
    elif section.key == "document_control":
        _render_document_control(doc, model, section)
    elif section.key == "table_of_contents":
        _render_table_of_contents(doc, model, section)
    else:
        # Should not happen if caller checks RESERVED_TEMPLATE_SECTIONS
        return


# -------------------------------------------------------------------
# Public API
# -------------------------------------------------------------------

def render_document(
    model: DocumentModel,
    template_path: str | Path,
    output_path: str | Path,
) -> None:
    """
    Render a DocumentModel into a Word document based on the given template.

    - The template is assumed to have:
        * Title page on page 1
        * Document Control on page 1 (with a metadata table)
        * Table of Contents on page 2
        * Headers/footers/page numbers/styles already configured

    - Reserved sections ('title_page', 'document_control', 'table_of_contents')
      do not have their layout rebuilt; they are handled by dedicated helpers.

    - All other sections are appended after the existing template content,
      with headings mapped to ISMS Heading styles and body content rendered
      according to ContentBlock.kind.

    Usage:
        model = DocumentModel.parse_file("examples/v2/sample_record.json")
        render_document(model, "templates/ISMS_Base_Master.docx", "out/REC-OPS-001.docx")
    """
    template_path = Path(template_path)
    output_path = Path(output_path)

    if not template_path.is_file():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    # First, handle reserved sections (title page, doc control, TOC)
    for section in model.sections:
        if section.key in RESERVED_TEMPLATE_SECTIONS:
            _dispatch_reserved_section(doc, model, section)

    # Apply [[DOC_*]] placeholders everywhere
    _apply_metadata_placeholders(doc, model.metadata)

    # Then render all non-reserved sections at the end of the document
    for section in model.sections:
        if section.key in RESERVED_TEMPLATE_SECTIONS:
            continue
        _render_section_recursive(doc, section)

    doc.save(str(output_path))
