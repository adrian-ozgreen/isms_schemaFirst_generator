"""
Helpers for adding content under existing headings in a Word template using python-docx.
Note: Custom Properties inside DOCX are not edited here (requires additional libs).
"""
from typing import List, Optional
from collections.abc import Sequence

from docx.document import Document as _Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def _find_heading_paragraph(doc: _Document, heading_text: str) -> Optional[Paragraph]:
    """Return the first paragraph whose text matches heading_text (case-insensitive, trimmed)."""
    target = (heading_text or "").strip().lower()
    if not target:
        return None
    for p in doc.paragraphs:
        if p.text.strip().lower() == target:
            return p
    return None


def _insert_after(paragraph: Paragraph, text: str = "", style_candidates=()) -> Paragraph:
    """
    Insert a new paragraph immediately AFTER the given paragraph, with optional style candidates.
    This is the key to ensuring content appears under the heading instead of at the end of the doc.
    """
    # Create a new <w:p> and insert it after the current paragraph's element
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)

    # Wrap it as a python-docx Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text or ""

    # Try the candidate styles in order
    for style_name in style_candidates or ():
        try:
            new_para.style = style_name
            break
        except KeyError:
            continue

    return new_para


def add_body_under_heading(doc: _Document, heading_text: str, body_text: str) -> None:
    """Append body text under a heading, creating the heading if needed with ISMS styles."""
    p = _find_heading_paragraph(doc, heading_text)
    if p is None:
        # Create a new heading with ISMS Heading 1 if available
        p = doc.add_paragraph(heading_text or "")
        for style_name in ("ISMS Heading 1", "Heading 1"):
            try:
                p.style = style_name
                break
            except KeyError:
                continue

    # Insert each line directly under the heading (or under the last inserted line)
    current = p
    for line in str(body_text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        current = _insert_after(
            current,
            line,
            style_candidates=("ISMS Body", "Normal"),
        )


def populate_revision_history(doc: _Document, rows: List[dict]) -> None:
    """Best-effort: find first table with 'Revision History' header and append rows."""
    for t in doc.tables:
        header = " ".join(cell.text.strip() for cell in t.rows[0].cells).lower()
        if "revision" in header and "version" in header and "date" in header:
            # Ensure TracWater table style
            try:
                t.style = "TracWater table"
            except Exception:
                pass

            for r in rows:
                row = t.add_row().cells
                # Expecting: Version | Date | Author | Changes Made | Approved By
                vals = [
                    r.get("version", ""),
                    r.get("date", ""),
                    r.get("author", ""),
                    r.get("changes", ""),
                    r.get("approved_by", ""),
                ]
                for i, v in enumerate(vals):
                    if i < len(row):
                        row[i].text = str(v)
            break


def add_numbered_list_under_heading(
    doc: _Document,
    heading_text: str,
    items: Sequence[str],
) -> None:
    """
    Append a numbered list under a heading, creating the heading if needed.

    Uses ISMS list styles if available, with sensible fallbacks.
    """
    # Ensure heading exists, same behaviour as add_body_under_heading
    p = _find_heading_paragraph(doc, heading_text)
    if p is None:
        p = doc.add_paragraph(heading_text or "")
        for style_name in ("ISMS Heading 1", "Heading 1"):
            try:
                p.style = style_name
                break
            except KeyError:
                continue

    current = p
    for item in items or []:
        text = str(item or "").strip()
        if not text:
            continue
        current = _insert_after(
            current,
            text,
            style_candidates=("ISMS List Number", "List Number", "ISMS Body", "Normal"),
        )


def add_bullet_list_under_heading(
    doc: _Document,
    heading_text: str,
    items: Sequence[str],
) -> None:
    """
    Append a bulleted list under a heading, creating the heading if needed.
    """
    p = _find_heading_paragraph(doc, heading_text)
    if p is None:
        p = doc.add_paragraph(heading_text or "")
        for style_name in ("ISMS Heading 1", "Heading 1"):
            try:
                p.style = style_name
                break
            except KeyError:
                continue

    current = p
    for item in items or []:
        text = str(item or "").strip()
        if not text:
            continue
        current = _insert_after(
            current,
            text,
            style_candidates=("ISMS List Bullet", "List Bullet", "ISMS Body", "Normal"),
        )
