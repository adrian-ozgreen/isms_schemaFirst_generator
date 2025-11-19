from typing import List, Dict, Any, Optional
from docx.document import Document as _Document
from docx.oxml.ns import qn

DEFAULT_TABLE_STYLE = "TracWater table"

def _find_paragraph(document: _Document, text: str):
    tgt = (text or "").strip().lower()
    for p in document.paragraphs:
        if (p.text or "").strip().lower() == tgt:
            return p
    return None

def _insert_table_after_paragraph(document: _Document, paragraph, rows: int, cols: int):
    # create at end, then move it directly after the paragraph (true placement)
    tbl = document.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(tbl._tbl)  # move table XML after the paragraph XML
    return tbl



def _iter_body_elements(document: _Document):
    body = document.element.body
    para_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = para_map.get(child)
            if p is not None:
                yield ("p", p)
        elif child.tag == qn('w:tbl'):
            t = table_map.get(child)
            if t is not None:
                yield ("t", t)


def _find_paragraph(document: _Document, text: str):
    tgt = (text or "").strip().lower()
    for p in document.paragraphs:
        if (p.text or "").strip().lower() == tgt:
            return p
    return None

def _insert_paragraph_after(document: _Document, paragraph, text: str, bold: bool = False, style: Optional[str] = None):
    # Create a paragraph at end, then move it to sit right after `paragraph`
    new_p = document.add_paragraph(text or "")
    if style:
        try:
            new_p.style = style
        except Exception:
            pass
    if bold:
        # Make all runs bold; if no runs exist yet, add one
        if new_p.runs:
            for r in new_p.runs:
                r.bold = True
        else:
            run = new_p.add_run(text or "")
            run.bold = True
            new_p.text = ""
    paragraph._p.addnext(new_p._p)
    return new_p

def _insert_table_after_paragraph(document: _Document, paragraph, rows: int, cols: int):
    # Create a table at end, then move it directly after `paragraph`
    tbl = document.add_table(rows=max(1, rows), cols=max(1, cols))
    paragraph._p.addnext(tbl._tbl)
    return tbl

def _get_table_cols(table) -> int:
    try:
        return len(table.rows[0].cells) if table.rows else 0
    except Exception:
        return 0




def _find_table_after_heading(document: _Document, heading_text: str):
    target = (heading_text or "").strip().lower()
    seen_heading = False
    for kind, el in _iter_body_elements(document):
        if kind == "p":
            if (el.text or "").strip().lower() == target:
                seen_heading = True
        elif kind == "t" and seen_heading:
            return el
    return None

def _insert_heading_and_table_after(document: _Document, heading_text: Optional[str]):
    # Insert an optional heading paragraph, then a minimal table at the end.
    if heading_text:
        document.add_paragraph(heading_text)
    t = document.add_table(rows=1, cols=1)
    return t

def _ensure_rows(table, count: int):
    while len(table.rows) < count:
        table.add_row()

def _set_cell_text(cell, text):
    cell.text = "" if text is None else str(text)

def populate_dynamic_table(doc: _Document, spec: Dict[str, Any]) -> bool:
    """
    JSON spec example:
    {
      "target": {"after_heading": "Equipment Register"},
      "create_if_missing": true,
      "heading": null,                       # optional: string to add as heading; null = no heading
      "heading_style": "Heading 2",          # reserved (not applied yet)
      "columns": ["Item","Serial","Installed At","Notes"],
      "rows": [
        ["PS3 Meter","SN-123","Main Switchboard",""],
        {"Item":"H2S Sensor","Installed At":"Blower Room","Notes":"Baseline offset applied"}
      ]
    }
    """
    target = spec.get("target", {}) or {}
    after_heading = target.get("after_heading")
    create_if_missing = bool(spec.get("create_if_missing", False))
    heading_text = spec.get("heading")
    columns: List[str] = spec.get("columns") or []
    rows = spec.get("rows") or []

    # Decide where to place a new table if needed
    insertion_anchor = None  # a paragraph after which we’ll insert heading and/or table

    if after_heading:
        anchor_para = _find_paragraph(doc, after_heading)
        if anchor_para is not None:
            insertion_anchor = anchor_para

    # Try to find an existing table immediately after the heading
    table = None
    if after_heading:
        table = _find_table_after_heading(doc, after_heading)

    # If no table found and creation is allowed, create in the right place
    if table is None and create_if_missing:
        # If we have an anchor and a heading text for the table, insert the heading paragraph first (bold body text)
        if insertion_anchor is not None:
            heading_para = None
            if spec.get("heading"):
                heading_para = _insert_paragraph_after(doc, insertion_anchor, spec["heading"], bold=True, style=None)
                insertion_anchor = heading_para  # next insert right after the heading
            # create the table after the (possibly new) anchor paragraph
            table = _insert_table_after_paragraph(
                doc, insertion_anchor, rows=1 + len(rows), cols=max(1, len(columns) or 1)
            )
        else:
            # No known anchor; append to end. If heading is requested, create a bold paragraph first.
            end_heading_para = None
            if spec.get("heading"):
                end_heading_para = _insert_paragraph_after(doc, doc.paragraphs[-1], spec["heading"], bold=True, style=None)
                table = _insert_table_after_paragraph(
                    doc, end_heading_para, rows=1 + len(rows), cols=max(1, len(columns) or 1)
                )
            else:
                table = doc.add_table(rows=1 + len(rows), cols=max(1, len(columns) or 1))

    if table is None:
        return False

    # Determine required grid size
    n_cols = max(len(columns), max((len(r) if isinstance(r, list) else len(columns)) for r in rows) if rows else 0)
    n_cols = max(n_cols, 1)
    n_rows = 1 + len(rows)  # header + data rows

    # If existing table has incorrect columns, replace with a fresh table in-place (preserve placement)
    existing_cols = _get_table_cols(table)
    if existing_cols != n_cols:
        # insert a new table right after the paragraph that precedes the current table, if we can infer it
        # For simplicity, append a new correctly-sized table right after the best known anchor (insertion_anchor), if available
        if insertion_anchor is not None:
            table = _insert_table_after_paragraph(doc, insertion_anchor, rows=n_rows, cols=n_cols)
        else:
            table = doc.add_table(rows=n_rows, cols=n_cols)
    else:
        # ensure enough rows
        while len(table.rows) < n_rows:
            table.add_row()

    # Style (default to "TracWater table" if none provided)
    table_style = spec.get("table_style") or DEFAULT_TABLE_STYLE
    try:
        table.style = table_style
    except Exception:
        pass

    # Header row
    for ci in range(n_cols):
        hdr = columns[ci] if ci < len(columns) else f"Col{ci+1}"
        table.rows[0].cells[ci].text = hdr

    # Data rows
    for ri, row in enumerate(rows, start=1):
        if isinstance(row, list):
            for ci in range(n_cols):
                table.rows[ri].cells[ci].text = (row[ci] if ci < len(row) else "") or ""
        elif isinstance(row, dict):
            for ci in range(n_cols):
                col_name = columns[ci] if ci < len(columns) else f"Col{ci+1}"
                table.rows[ri].cells[ci].text = str(row.get(col_name, "") or "")
        else:
            table.rows[ri].cells[0].text = str(row)




def _get_table_cols(table) -> int:
    try:
        return len(table.rows[0].cells) if table.rows else 0
    except Exception:
        return 0
