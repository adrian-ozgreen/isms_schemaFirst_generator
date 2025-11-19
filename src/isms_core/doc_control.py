
from docx.document import Document as _Document
from docx.oxml.ns import qn

DOC_CONTROL_LABELS = {
    "Document ID": "doc_id",
    "Title": "title",
    "Version": "version",
    "Owner": "owner",
    "Approved By": "approved_by",
    "Confidentiality": "confidentiality",
    "Status": "status",
    "Date Completed": "date_completed",
    "Next Review Date": "next_review_date",
    "Document Type": "document_type",
}

def _iter_body_elements(document: _Document):
    """Yield paragraphs and tables in document order."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            # paragraph
            for p in document.paragraphs:
                if p._p is child:
                    yield p
                    break
        elif child.tag == qn('w:tbl'):
            # table
            for t in document.tables:
                if t._tbl is child:
                    yield t
                    break

def _find_table_after_heading(document: _Document, heading_text: str):
    target = (heading_text or "").strip().lower()
    seen_heading = False
    for el in _iter_body_elements(document):
        if hasattr(el, "text"):  # paragraph
            if (el.text or "").strip().lower() == target:
                seen_heading = True
        else:
            # table
            if seen_heading:
                return el
    return None

def _find_table_with_cell_text(document: _Document, needle: str):
    n = (needle or "").strip().lower()
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                if n in (cell.text or "").strip().lower():
                    return t
    return None

def _ensure_two_columns(t):
    # Best-effort: ensure two columns, if fewer than 2, add a second cell to each row.
    for r in t.rows:
        if len(r.cells) < 2:
            r.add_cell()

def populate_document_control_table(doc: _Document, meta: dict) -> bool:
    """
    Locate the Document Control table and populate it using metadata.
    Strategy:
      1) Find table immediately after a paragraph with exact text 'Document Control'
      2) Fallback: find any table containing a cell with 'Document Control'
    """
    table = _find_table_after_heading(doc, "Document Control")
    if table is None:
        table = _find_table_with_cell_text(doc, "Document Control")
    if table is None:
        return False

    _ensure_two_columns(table)

    # Build index of existing labels
    existing = {}
    for i, row in enumerate(table.rows):
        if len(row.cells) >= 2:
            key = (row.cells[0].text or "").strip()
            if key:
                existing[key] = i

    # Ensure each property row exists and set values
    def ensure_row(label: str) -> int:
        if label in existing:
            return existing[label]
        row = table.add_row()
        row.cells[0].text = label
        return len(table.rows) - 1

    for label, meta_key in DOC_CONTROL_LABELS.items():
        idx = ensure_row(label)
        try:
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = str(meta.get(meta_key, "") or "")
        except Exception:
            pass

    return True
